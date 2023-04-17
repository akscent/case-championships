import io
import os
import shutil
from unidecode import unidecode
import chardet as chardet
from fastapi import HTTPException
from docx import Document
from git import Repo



class Git:
    def __init__(self):
        self.repo_path = os.path.join(os.getcwd(), 'word_docs')
        if os.path.exists(self.repo_path):
            self.repo = Repo.init(self.repo_path)
        else:
            self.repo = Repo(self.repo_path)

    def new_commit(self, file_name, commit_name, file,):
        doc_path = os.path.join(self.repo_path, file_name)
        if not(os.path.exists(doc_path)):
            raise HTTPException(status_code=404, detail="Current fail doesn't exist")
        document = Document(io.BytesIO(file))
        document.save(doc_path)
        index = self.repo.index
        index.add([doc_path])
        index.commit(commit_name)

        return "ok"

    def get_commits(self, file_name):
        commits = []
        for commit in self.repo.iter_commits(paths=file_name):
            commits.append(commit.message)

        return commits

    def get_files(self):
        return os.listdir("word_docs")

    def get_commit(self, commit_name, file_name):
        for commit in self.repo.iter_commits(paths=file_name):
            if commit.message == commit_name:
                contents = commit.tree[file_name].data_stream.read()
                document = Document(io.BytesIO(contents))
                paragraphs = []
                for paragraph in document.paragraphs:
                    paragraphs.append(paragraph.text)

                return paragraphs
